import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    if level == 1:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        return p
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = False
        run.font.size = Pt(14)
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.space_before = Pt(12)
        return p

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.line_spacing = 1.5
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def main():
    doc = Document()
    
    # Set default style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.right_margin = Inches(1)
        section.left_margin = Inches(1.5)

    # ---------------- TITLE PAGE ----------------
    
    def add_centered_line(text, size_pt, bold=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size_pt)
        run.bold = bold
        p.paragraph_format.line_spacing = 1.5
        return p

    doc.add_paragraph() # Top padding
    add_centered_line("Synopsis", 14)
    add_centered_line("on", 14)
    p_title = add_centered_line("SHINETWOPLAY", 24, bold=True)
    p_title.paragraph_format.space_before = Pt(24)
    p_title.paragraph_format.space_after = Pt(24)
    
    add_centered_line("SUBMITTED TOWARDS PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE AWARD OF THE DEGREE OF", 12)
    
    p = add_centered_line("BACHELOR OF TECHNOLOGY", 14, bold=True)
    p.paragraph_format.space_before = Pt(12)
    add_centered_line("(Computer Science & Engineering)", 14)
    
    p = add_centered_line("SUBMITTED BY", 12)
    p.paragraph_format.space_before = Pt(24)
    
    add_centered_line("Soumya (Roll No. 2211031)", 14, bold=True)
    add_centered_line("Nikhil (Roll No. 2211066)", 14, bold=True)
    
    p = add_centered_line("Under the supervision of", 12)
    p.paragraph_format.space_before = Pt(24)
    
    add_centered_line("NAME OF THE SUPERVISOR", 14, bold=True)
    add_centered_line("(Department of Computer Science & Engineering)", 12)
    
    p = add_centered_line("[ Logo of MDU ]", 14)
    p.paragraph_format.space_before = Pt(36)
    p.paragraph_format.space_after = Pt(36)
    
    add_centered_line("University Institute of Engineering and Technology", 14, bold=True)
    add_centered_line("(Maharshi Dayanand University, Rohtak)", 14)
    
    p = add_centered_line("Jan-June, 2026", 12)
    p.paragraph_format.space_before = Pt(24)
    
    doc.add_page_break()
    
    # ---------------- CONTENT ----------------
    
    add_heading(doc, "1. Title", level=1)
    add_paragraph(doc, "ShineTwoPlay: A Real-Time Multiplayer Gaming and Communication Platform")
    
    add_heading(doc, "2. Introduction", level=1)
    add_paragraph(doc, "ShineTwoPlay is a real-time multiplayer gaming platform designed exclusively for two players. It aims to provide a frictionless entertainment experience where connection, creativity, and fun come together seamlessly. Unlike traditional platforms that require heavy downloads or lengthy sign-up processes, ShineTwoPlay operates entirely within the browser. It features 16 diverse multiplayer games, integrated premium chat, and WebRTC-powered voice calling, bridging the gap between casual gaming and real-time communication.")
    
    add_heading(doc, "3. Objectives", level=1)
    p = add_paragraph(doc, "")
    p.add_run("• To develop a highly responsive, browser-based multiplayer gaming platform that requires no installation or registration.\n")
    p.add_run("• To implement a diverse catalog of 16 real-time and turn-based games (e.g., Ludo, Tic Tac Toe, Paddle Arena).\n")
    p.add_run("• To integrate premium communication features, including WebRTC voice calling, voice messages, and image sharing, directly into the gaming environment.\n")
    p.add_run("• To ensure seamless, low-latency synchronization of game state between connected clients using WebSockets and Redis.")

    add_heading(doc, "4. Scope of the Project", level=1)
    add_paragraph(doc, "The scope of ShineTwoPlay is focused on delivering a 2-player interactive experience over the web. The platform boundaries are limited to a maximum of two players per room, functioning via dynamically generated 4-letter room codes. Persistent user accounts and global scale matchmaking are outside the current scope, as the project prioritizes fast, transient, and invite-only gaming sessions. The system relies heavily on modern web browser capabilities like WebSockets and WebRTC.")

    add_heading(doc, "5. Literature Review", level=1)
    add_paragraph(doc, "A review of existing literature and current market offerings reveals that while casual web games are abundant, they often lack integrated real-time communication tools, forcing users to rely on third-party applications (like Discord or Skype) while playing. Conversely, traditional multiplayer games often pose high entry barriers with mandatory sign-ups and downloads. Academic studies on WebSocket and WebRTC technologies highlight their capacity for sub-second latency and peer-to-peer data exchange in modern browsers. ShineTwoPlay bridges the identified gap by combining instant-access web gaming with built-in voice and text communication using these advanced protocols.")

    add_heading(doc, "6. Proposed System/Implementation", level=1)
    add_paragraph(doc, "The proposed system follows a client-server architecture optimized for real-time bidirectional communication. The backend is powered by Django and Django Channels, utilizing Redis as the channel layer to broadcast messages instantly between players in the same room. The frontend features a modern 'Glassmorphism' user interface built with HTML5, CSS3, and embedded JavaScript. For voice communication, the platform implements WebRTC to establish direct peer-to-peer audio connections. The game state is transiently managed in memory and synchronized via custom WebSocket events for each of the 16 distinct games.")

    add_heading(doc, "7. Tools and Technologies", level=1)
    p = add_paragraph(doc, "")
    p.add_run("- Backend Framework: Django 4.2\n")
    p.add_run("- Real-time Handling: Django Channels 4.0, ASGI (Daphne)\n")
    p.add_run("- In-Memory Data Store / Broker: Redis\n")
    p.add_run("- Database: SQLite\n")
    p.add_run("- Frontend: HTML5, CSS3, Vanilla JavaScript\n")
    p.add_run("- Communication Protocols: WebSockets, WebRTC\n")
    p.add_run("- Deployment: Nginx, Certbot SSL, Ubuntu/EC2")

    add_heading(doc, "8. Project Timeline", level=1)
    add_paragraph(doc, "(Note: A flowchart representation of the timeline will be included in the final report. Below are the key milestones.)")
    p = add_paragraph(doc, "")
    p.add_run("- January 2026: Requirement Analysis, System Architecture Design, and Frontend UI/UX (Glassmorphism) Prototyping.\n")
    p.add_run("- February 2026: Backend Setup, Django Channels Integration, and WebSocket room management implementation.\n")
    p.add_run("- March 2026: Development and integration of the core 16 multiplayer games (logic and synchronization).\n")
    p.add_run("- April 2026: Implementation of premium communication features (Chat, WebRTC Voice Calls, Image/Voice Media sharing).\n")
    p.add_run("- May 2026: Comprehensive System Testing, Bug Fixing, and Performance Optimization.\n")
    p.add_run("- June 2026: Final Project Deployment, security hardening, and Documentation compilation.")

    add_heading(doc, "9. Expected Outcome", level=1)
    add_paragraph(doc, "The expected outcome is a fully functional, scalable, and responsive web application where two users can instantly join a shared room, engage in voice and text communication, and play seamlessly synchronized multiplayer games without any perceivable lag or requirement for account registration.")

    add_heading(doc, "10. Significance of the Project", level=1)
    add_paragraph(doc, "ShineTwoPlay significantly lowers the barrier to entry for shared online entertainment. It provides a zero-friction platform for friends to connect and interact. Technically, it serves as a robust demonstration of the capabilities of modern web standards, illustrating how Django Channels, WebSockets, and WebRTC can be orchestrated to build complex, low-latency, stateful web applications.")

    add_heading(doc, "11. Future Scope", level=1)
    add_paragraph(doc, "Future enhancements could include implementing an AI opponent system to allow for single-player modes when a partner is unavailable. Additionally, the platform could be expanded to support larger multiplayer rooms, persistent leaderboards, user profiles using secure authentication, and a progressive web app (PWA) version for native-like mobile experiences.")

    add_heading(doc, "12. Conclusion", level=1)
    add_paragraph(doc, "In conclusion, ShineTwoPlay is an innovative approach to online casual gaming. By stripping away accounts and downloads, and directly embedding high-quality communication tools alongside engaging gameplay, the project creates a unique, accessible, and highly connected digital experience for two players.")

    add_heading(doc, "13. References", level=1)
    p = add_paragraph(doc, "")
    p.add_run("1. Grigorik, I. (2013). High Performance Browser Networking. O'Reilly Media.\n")
    p.add_run("2. Lubbers, P., Albers, B., & Salim, V. (2010). Pro HTML5 Programming: Powerful APIs for Richer Internet Application Development. Apress.\n")
    p.add_run("3. Johnston, A. B., & Burnett, D. C. (2012). WebRTC: APIs and RTCWEB Protocols of the HTML5 Web.\n")
    p.add_run("4. Kurose, J. F., & Ross, K. W. (2017). Computer Networking: A Top-Down Approach (7th ed.). Pearson.\n")
    p.add_run("5. Holovaty, A., & Kaplan-Moss, J. (2009). The Definitive Guide to Django: Web Development Done Right. Apress.\n")
    p.add_run("6. Freeman, E., & Robson, E. (2014). Head First JavaScript Programming. O'Reilly Media.\n")
    p.add_run("7. Flanagan, D. (2020). JavaScript: The Definitive Guide (7th ed.). O'Reilly Media.\n")
    p.add_run("8. Coulouris, G., Dollimore, J., Kindberg, T., & Blair, G. (2011). Distributed Systems: Concepts and Design (5th ed.). Addison-Wesley.\n")
    p.add_run("9. Pressman, R. S. (2014). Software Engineering: A Practitioner's Approach (8th ed.). McGraw-Hill Education.\n")
    p.add_run("10. Sommerville, I. (2015). Software Engineering (10th ed.). Pearson.")

    # Apply times new roman to all runs
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if not run.font.name:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

    doc.save('ShineTwoPlay_Synopsis.docx')
    print("Synopsis Word document created successfully as ShineTwoPlay_Synopsis.docx")

if __name__ == '__main__':
    main()
